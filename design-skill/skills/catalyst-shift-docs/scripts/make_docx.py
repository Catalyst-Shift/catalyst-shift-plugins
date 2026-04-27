"""
Catalyst Shift — Markdown → DOCX converter
===========================================
Takes a markdown file (e.g. an internal doc, or a flattened proposal/SOW)
and emits a Word doc styled with the Catalyst Shift type stack.

Headings, body, and tables get CS-aligned fonts and colors. Not pixel-
perfect to the HTML templates — use this when the teammate needs an
*editable* Word doc rather than a print-faithful PDF.

Requirements:
    pip install python-docx markdown beautifulsoup4

Usage:
    python make_docx.py input.md output.docx [--title "Doc Title"]
"""

import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
    from bs4 import BeautifulSoup, NavigableString
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Install with: pip install python-docx markdown beautifulsoup4")
    sys.exit(1)


# Catalyst Shift tokens
GOLD = RGBColor(0xC8, 0xA9, 0x6E)
TEAL = RGBColor(0x2D, 0xD4, 0xA8)
TEXT_PRIMARY = RGBColor(0x1A, 0x1A, 0x1F)
TEXT_SECONDARY = RGBColor(0x5C, 0x58, 0x50)
TEXT_MUTED = RGBColor(0x8C, 0x87, 0x80)

DISPLAY_FONT = "Calibri"   # General Sans not always installed; Calibri reads cleanest
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def style_doc(doc: Document, title: str | None) -> None:
    """Set base styles to match the CS feel as close as Word allows."""
    for s in doc.styles:
        if s.name == "Normal":
            s.font.name = BODY_FONT
            s.font.size = Pt(11)
            s.font.color.rgb = TEXT_SECONDARY
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    if title:
        p = doc.add_paragraph()
        run = p.add_run("CATALYST SHIFT")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = GOLD
        run.bold = True
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.font.name = DISPLAY_FONT
        hr.font.size = Pt(28)
        hr.font.color.rgb = TEXT_PRIMARY
        hr.bold = True
        doc.add_paragraph()  # spacer


def add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 22, 2: 18, 3: 14, 4: 11}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    if level <= 2:
        kicker = p.add_run(f"§ ")
        kicker.font.name = MONO_FONT
        kicker.font.size = Pt(9)
        kicker.font.color.rgb = GOLD
        kicker.bold = True
    run = p.add_run(text)
    run.font.name = DISPLAY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_PRIMARY
    run.bold = True


def add_paragraph(doc: Document, soup_p) -> None:
    p = doc.add_paragraph()
    for child in soup_p.children:
        if isinstance(child, NavigableString):
            run = p.add_run(str(child))
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_SECONDARY
        elif child.name == "strong":
            run = p.add_run(child.get_text())
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_PRIMARY
            run.bold = True
        elif child.name == "em":
            run = p.add_run(child.get_text())
            run.italic = True
            run.font.color.rgb = GOLD
        elif child.name == "code":
            run = p.add_run(child.get_text())
            run.font.name = MONO_FONT
            run.font.color.rgb = TEAL
        else:
            run = p.add_run(child.get_text())
            run.font.name = BODY_FONT
            run.font.color.rgb = TEXT_SECONDARY


def add_list(doc: Document, soup_list, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in soup_list.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(li.get_text())
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_SECONDARY


def add_table(doc: Document, soup_table) -> None:
    rows = soup_table.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["td", "th"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        is_head = tr.find("th") is not None
        for j, cell in enumerate(cells):
            tc = table.rows[i].cells[j]
            tc.text = ""
            p = tc.paragraphs[0]
            run = p.add_run(cell.get_text())
            run.font.name = MONO_FONT if is_head else BODY_FONT
            run.font.size = Pt(9 if is_head else 10)
            run.font.color.rgb = TEXT_MUTED if is_head else TEXT_PRIMARY
            run.bold = is_head


def render(md_text: str, out_path: Path, title: str | None) -> None:
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    style_doc(doc, title)
    for el in soup.children:
        if not getattr(el, "name", None):
            continue
        if el.name in {"h1", "h2", "h3", "h4"}:
            add_heading(doc, el.get_text(), int(el.name[1]))
        elif el.name == "p":
            add_paragraph(doc, el)
        elif el.name == "ul":
            add_list(doc, el, ordered=False)
        elif el.name == "ol":
            add_list(doc, el, ordered=True)
        elif el.name == "table":
            add_table(doc, el)
        elif el.name == "hr":
            doc.add_paragraph("―" * 30)
        elif el.name == "blockquote":
            for p_ in el.find_all("p"):
                add_paragraph(doc, p_)
    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--title", default=None)
    args = parser.parse_args()
    md = Path(args.input).read_text(encoding="utf-8")
    render(md, Path(args.output), args.title)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
